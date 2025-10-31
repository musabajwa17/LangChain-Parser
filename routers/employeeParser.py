# import os
# import json
# import re
# from fastapi import APIRouter, UploadFile, File
# from fastapi.responses import JSONResponse
# from langchain_groq import ChatGroq
# from langchain.prompts import PromptTemplate
# from PyPDF2 import PdfReader
# from pdf2image import convert_from_path
# from docx import Document
# import pytesseract
# from dotenv import load_dotenv

# # ----------------------------
# # Load environment variables
# # ----------------------------
# load_dotenv()

# router = APIRouter()

# # ----------------------------
# # Initialize LLM
# # ----------------------------
# llm = ChatGroq(
#     model="llama-3.1-8b-instant",
#     temperature=0,
#     groq_api_key=os.getenv("GROQ_API_KEY")
# )

# # ----------------------------
# # PDF TEXT EXTRACTION (with OCR fallback)
# # ----------------------------
# def extract_text_from_pdf(file_path: str) -> str:
#     text = ""
#     try:
#         reader = PdfReader(file_path)
#         for page in reader.pages:
#             page_text = page.extract_text()
#             if page_text:
#                 text += page_text
#     except Exception:
#         pass

#     if not text.strip():
#         try:
#             images = convert_from_path(file_path)
#             ocr_text = "\n".join(pytesseract.image_to_string(img) for img in images)
#             text += ocr_text
#         except Exception as e:
#             print("OCR failed:", e)
#     return text.strip()


# # ----------------------------
# # DOCX TEXT EXTRACTION
# # ----------------------------
# def extract_text_from_docx(file_path: str) -> str:
#     try:
#         doc = Document(file_path)
#         return "\n".join([para.text for para in doc.paragraphs])
#     except Exception as e:
#         print("DOCX extraction failed:", e)
#         return ""


# # ----------------------------
# # CLEAN JSON OUTPUT
# # ----------------------------
# def clean_json_output(text: str):
#     text = text.replace("```json", "").replace("```", "").strip()
#     text = re.sub(r'^[^{]*', '', text)
#     text = re.sub(r'[^}]*$', '', text)
#     try:
#         return json.loads(text)
#     except json.JSONDecodeError:
#         return {"error": "Invalid JSON output from LLM", "raw_output": text}


# # ----------------------------
# # PROMPT TEMPLATE
# # ----------------------------
# template = PromptTemplate(
#     input_variables=["resume_text"],
#     template="""
# Extract the following structured information from the resume below:

# - name
# - email
# - phone
# - citations
# - impactFactor
# - scholar
# - education (degree, institution, year)
# - experience (role, company, years)
# - ACHIEVEMENTS
# - BOOKAUTHORSHIP
# - journalGuestEditor
# - RESEARCHPUBLICATIONS (journal,workshop)
# - mssupervised (studentName, thesisTitle, year)
# - phdstudentsupervised (studentName, thesisTitle, year)
# - RESEARCHPROJECTS (title,description)
# - PROFESSIONAACTIVITIES
# - PROFESSIONALTRAINING (title, description,year)
# - technicalSkills (category, details)  
# - MEMBERSHIPS&OTHERASSOCIATIONS (heading, desc)
# - REFERENCE (prof, designation,mail,phone)

# Return ONLY valid JSON format (no explanations, no extra text).

# Resume Text:
# {resume_text}
# """
# )

# # ----------------------------
# # PARSING ENDPOINT
# # ----------------------------
# @router.post("/employee-parser")
# async def parse_resume(file: UploadFile = File(...)):
#     try:
#         if not (file.filename.lower().endswith(".pdf") or file.filename.lower().endswith(".docx")):
#             return JSONResponse(content={"error": "Unsupported file type. Please upload PDF or DOCX only."}, status_code=400)

#         temp_path = f"temp_{file.filename}"
#         with open(temp_path, "wb") as f:
#             f.write(await file.read())

#         if file.filename.lower().endswith(".pdf"):
#             resume_text = extract_text_from_pdf(temp_path)
#         else:
#             resume_text = extract_text_from_docx(temp_path)

#         if not resume_text.strip():
#             os.remove(temp_path)
#             return JSONResponse(content={"error": "No readable text found. Try uploading a text-based resume."}, status_code=400)

#         resume_text = resume_text[:15000]
#         chain = template | llm
#         structured_response = chain.invoke({"resume_text": resume_text}).content
#         structured_data = clean_json_output(structured_response)

#         os.remove(temp_path)
#         return JSONResponse(content=structured_data)

#     except Exception as e:
#         return JSONResponse(content={"error": str(e)}, status_code=500)


import os
import json
import re
import ast
from fastapi import APIRouter, UploadFile, File
from fastapi.responses import JSONResponse
from langchain_groq import ChatGroq
from langchain.prompts import PromptTemplate
from PyPDF2 import PdfReader
from pdf2image import convert_from_path
from docx import Document
import pytesseract
from dotenv import load_dotenv

# ----------------------------
# Load environment variables
# ----------------------------
load_dotenv()
router = APIRouter()

# ----------------------------
# Initialize LLM
# ----------------------------
llm = ChatGroq(
    model="llama-3.1-8b-instant",
    temperature=0,
    groq_api_key=os.getenv("GROQ_API_KEY")
)

# ----------------------------
# TEXT EXTRACTION HELPERS
# ----------------------------
def extract_text_from_pdf(file_path: str) -> str:
    text = ""
    try:
        reader = PdfReader(file_path)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text
    except Exception:
        pass

    # OCR fallback if PDF text extraction fails
    if not text.strip():
        try:
            images = convert_from_path(file_path)
            ocr_text = "\n".join(pytesseract.image_to_string(img) for img in images)
            text += ocr_text
        except Exception as e:
            print("⚠️ OCR failed:", e)
    return text.strip()


def extract_text_from_docx(file_path: str) -> str:
    """Extracts text from DOCX resumes."""
    try:
        doc = Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        print("⚠️ DOCX extraction failed:", e)
        return ""


# ----------------------------
# CLEAN & FIX JSON OUTPUT
# ----------------------------
def clean_json_output(text: str):
    """
    Cleans and normalizes LLM output into valid JSON.
    Handles malformed structures, smart quotes, and trailing commas.
    """
    try:
        text = text.replace("```json", "").replace("```", "").strip()
        text = re.sub(r'^[^{]*', '', text)
        text = re.sub(r'[^}]*$', '', text)
        text = (
            text.replace("“", '"')
                .replace("”", '"')
                .replace("’", "'")
                .replace("\n", " ")
                .replace("\r", " ")
        )
        text = re.sub(r",\s*}", "}", text)
        text = re.sub(r",\s*]", "]", text)
        return json.loads(text)
    except json.JSONDecodeError as e:
        print("⚠️ JSON decode error:", e)
        try:
            return ast.literal_eval(text)
        except Exception as inner_e:
            print("⚠️ Fallback parse failed:", inner_e)
            return {"error": "Invalid JSON output from LLM", "raw_output": text}


# ----------------------------
# SECTION PREPROCESSORS
# ----------------------------
def remove_research_publications(text: str) -> str:
    """
    Removes 'Research Publications', 'Books', or 'Conferences' sections
    to prevent confusion with supervision sections.
    """
    clean_text = re.sub(r'\s+', ' ', text)
    pattern = r'(RESEARCH\s*PUBLICATIONS.*?)(M\.?S\.?\s*STUDENTS?\s*SUPERVISED|PhD\s*STUDENTS?\s*SUPERVISED|$)'
    result = re.sub(pattern, r'\2', clean_text, flags=re.IGNORECASE)
    return result.strip()


def extract_supervision_sections(text: str):
    """
    Extracts M.S. and PhD supervision sections separately for better accuracy.
    """
    ms_section = ""
    phd_section = ""

    clean_text = re.sub(r'\s+', ' ', text)
    clean_text = (
        clean_text.replace("Ph.D.", "PhD")
        .replace("M. S.", "M.S.")
        .replace("M S", "M.S.")
    )

    ms_match = re.search(
        r'(M\.?S\.?\s*STUDENTS?\s*SUPERVISED.*?)(PhD\s*STUDENTS?\s*SUPERVISED|$)',
        clean_text,
        re.IGNORECASE,
    )
    if ms_match:
        ms_section = ms_match.group(1).strip()

    phd_match = re.search(
        r'(PhD\s*STUDENTS?\s*SUPERVISED.*)',
        clean_text,
        re.IGNORECASE,
    )
    if phd_match:
        phd_section = phd_match.group(1).strip()

    return ms_section, phd_section


# ----------------------------
# PROMPT TEMPLATE (STRICT + GUIDED)
# ----------------------------
template = PromptTemplate(
    input_variables=["resume_text"],
    template="""
Extract structured information from the resume text below.

Return **only valid JSON** with the following structure:

{{
  "name": "",
  "email": "",
  "phone": "",
  "citations": "",
  "impactFactor": "",
  "scholar": "",
  "education": [{{"degree": "", "institution": "", "year": ""}}],
  "experience": [{{"role": "", "company": "", "years": ""}}],
  "achievements": [""],
  "bookAuthorship": [{{"title": "", "publisher": ""}}],
  "journalGuestEditor": [{{"title": "", "publisher": "", "section": ""}}],
  "researchPublications": [{{"title": "", "journal": "", "year": ""}}],
  "mssupervised": [{{"studentName": "", "thesisTitle": "", "year": ""}}],
  "phdstudentsupervised": [{{"studentName": "", "thesisTitle": "", "year": ""}}],
  "researchProjects": [{{"title": "", "description": ""}}],
  "professionalActivities": [{{"heading": "", "desc": "", "year": ""}}],
  "professionalTraining": [{{"title": "", "description": "", "year": ""}}],
  "technicalSkills": [{{"category": "", "details": ""}}],
  "membershipsAndOtherAssociations": [{{"heading": "", "desc": "", "year": ""}}],
  "reference": [{{"prof": "", "designation": "", "mail": "", "phone": ""}}]
}}

Rules:
1. Ignore any research publication titles when listing student supervision.
2. Use the text markers “--- START OF M.S. SUPERVISED SECTION ---” and “--- START OF PhD SUPERVISED SECTION ---” for accurate extraction.
3. Do not duplicate students between M.S. and PhD arrays.
4. Only include meaningful fields (ignore null/empty).
5. Return **only JSON**, with no explanations or markdown.

Resume Text:
{resume_text}
"""
)

# ----------------------------
# PARSING ENDPOINT
# ----------------------------
@router.post("/employee-parser")
async def parse_resume(file: UploadFile = File(...)):
    try:
        # Validate file type
        if not (file.filename.lower().endswith(".pdf") or file.filename.lower().endswith(".docx")):
            return JSONResponse(
                content={"error": "Unsupported file type. Please upload PDF or DOCX only."},
                status_code=400,
            )

        temp_path = f"temp_{file.filename}"
        with open(temp_path, "wb") as f:
            f.write(await file.read())

        # Extract resume text
        if file.filename.lower().endswith(".pdf"):
            resume_text = extract_text_from_pdf(temp_path)
        else:
            resume_text = extract_text_from_docx(temp_path)

        if not resume_text.strip():
            os.remove(temp_path)
            return JSONResponse(
                content={"error": "No readable text found. Try uploading a text-based resume."},
                status_code=400,
            )

        # Limit & preprocess text
        resume_text = resume_text[:15000]
        resume_text = remove_research_publications(resume_text)
        ms_text, phd_text = extract_supervision_sections(resume_text)

        # Add section markers for clarity
        resume_input = f"""
{resume_text}

--- START OF M.S. SUPERVISED SECTION ---
{ms_text}

--- START OF PhD SUPERVISED SECTION ---
{phd_text}
"""

        # Run the model
        chain = template | llm
        structured_response = chain.invoke({"resume_text": resume_input}).content
        print("🧩 Raw LLM output preview:", structured_response[:300])

        structured_data = clean_json_output(structured_response)

        os.remove(temp_path)
        return JSONResponse(content=structured_data)

    except Exception as e:
        print("❌ Unexpected Error:", e)
        return JSONResponse(content={"error": str(e)}, status_code=500)
